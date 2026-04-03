"""
Export docx_testset.jsonl to a readable Word document.
Usage: uv run python export_testset_to_docx.py
Output: docx_testset_review.docx
"""

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT_FILE  = Path(__file__).parent / "testsets" / "docx_testset.jsonl"
OUTPUT_FILE = Path(__file__).parent / "results" / "docx_testset_review.docx"


def add_styled_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_label_paragraph(doc, label, color_rgb=(0x1F, 0x49, 0x7D)):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor(*color_rgb)
    run.font.size = Pt(11)
    return p


def add_body_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.2)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


def main():
    records = []
    with open(INPUT_FILE, encoding="utf-8") as f:
        for line in f:
            if line.strip():
                records.append(json.loads(line))

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Title
    title = doc.add_heading("CRDCDH Chatbot — Q&A Review", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"{len(records)} questions generated from Bedrock Knowledge Base"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # spacer

    for i, rec in enumerate(records, 1):
        question     = rec.get("question", "")
        ground_truth = rec.get("ground_truth", "")
        metadata     = rec.get("metadata", {})
        n_chunks     = len(rec.get("contexts", []))
        qtype        = metadata.get("question_type", "")

        # Question heading
        add_styled_heading(doc, f"Q{i}. {question}", level=2)

        # Metadata line
        meta_p = doc.add_paragraph()
        meta_run = meta_p.add_run(f"Type: {qtype}   |   Retrieved chunks: {n_chunks}")
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        meta_run.italic = True

        # Answer label
        add_label_paragraph(doc, "Answer:")

        # Answer body — split on newlines to preserve structure
        for line in ground_truth.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            # Render markdown-style headers as bold lines
            if stripped.startswith("### "):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                r = p.add_run(stripped.lstrip("# "))
                r.bold = True
                r.font.size = Pt(10.5)
            elif stripped.startswith("## "):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                r = p.add_run(stripped.lstrip("# "))
                r.bold = True
                r.font.size = Pt(11)
            elif stripped.startswith("**") and stripped.endswith("**"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                r = p.add_run(stripped.strip("*"))
                r.bold = True
                r.font.size = Pt(10.5)
            else:
                add_body_paragraph(doc, stripped)

        # Divider
        doc.add_paragraph("─" * 80).paragraph_format.space_after = Pt(6)

    doc.save(OUTPUT_FILE)
    print(f"Saved {OUTPUT_FILE}  ({len(records)} Q&A pairs)")


if __name__ == "__main__":
    main()
